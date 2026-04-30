"""
DOCX parser: converts a CV .docx into the standard JSON schema.

Mirrors the PDF parser API:
    extract_docx_text(path) -> str
    parse_docx_to_cv(path, client, output_path=None) -> dict
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from .pdf_parser import _LLMClient, text_to_cv_json


def extract_docx_text(docx_path: str | Path) -> str:
    """Extract plain text from a .docx — paragraphs + tables."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install it with: pip install python-docx"
        ) from e

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX does not exist: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"File does not look like a .docx: {docx_path}")

    doc = Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ValueError(
            f"Could not extract text from DOCX {docx_path}. "
            "It may be empty or only contain images."
        )
    return full_text


def parse_docx_to_cv(
    docx_path: str | Path,
    client: _LLMClient,
    output_path: str | Path | None = None,
) -> dict[str, Any]:
    """Full pipeline: DOCX → CV JSON dict."""
    text = extract_docx_text(docx_path)
    cv_dict = text_to_cv_json(text, client)
    if output_path is not None:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(
            json.dumps(cv_dict, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    return cv_dict
