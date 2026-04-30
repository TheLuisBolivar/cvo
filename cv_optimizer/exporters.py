"""
Export the optimized CV to multiple formats: Markdown, JSON, PDF, DOCX.

- Markdown: written by `generator.generate_markdown` (already exists).
- JSON:     structured optimized CV (built by `generator.build_optimized_cv_dict`).
- PDF:      via `pandoc` (must be installed). Falls back with a clear warning.
- DOCX:     via `python-docx` (declared as a dependency).

`export_all` writes each requested format and returns a dict of
{format: written_path}. The Markdown is always the source of truth — the
DOCX is built from the same content the MD has.
"""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path
from typing import Any

from .models import CV


SUPPORTED_FORMATS = ("md", "json", "pdf", "docx")


# ──────────────────────────────────────────────────────────────────────
# Format selection
# ──────────────────────────────────────────────────────────────────────
def parse_format_list(raw: str | None) -> list[str]:
    """
    Parse a `--format` value (comma-separated, or "all"). Defaults to ["md"].
    Always includes "json" as a free side-effect because the structured CV
    JSON costs nothing to write and is useful downstream.
    """
    if not raw:
        return ["md", "json"]
    raw = raw.strip().lower()
    if raw == "all":
        return list(SUPPORTED_FORMATS)
    parts = [p.strip() for p in raw.split(",") if p.strip()]
    unknown = [p for p in parts if p not in SUPPORTED_FORMATS]
    if unknown:
        raise ValueError(
            f"Unsupported format(s): {', '.join(unknown)}. "
            f"Choose from: {', '.join(SUPPORTED_FORMATS)}, or 'all'."
        )
    if "json" not in parts:
        parts.append("json")
    return parts


# ──────────────────────────────────────────────────────────────────────
# Per-format writers
# ──────────────────────────────────────────────────────────────────────
def write_markdown(md: str, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(md, encoding="utf-8")
    return path


def write_json(cv_dict: dict[str, Any], path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(cv_dict, ensure_ascii=False, indent=2, default=_json_default),
        encoding="utf-8",
    )
    return path


def _json_default(o: Any) -> Any:
    """Fallback for dataclass values like Experience that may sneak in."""
    if hasattr(o, "__dict__"):
        return o.__dict__
    return str(o)


def write_pdf(md_path: Path, path: Path) -> Path:
    """
    Convert the Markdown CV to PDF via pandoc. Raises RuntimeError if pandoc
    is not on PATH so the caller can surface a friendly message.
    """
    if shutil.which("pandoc") is None:
        raise RuntimeError(
            "pandoc not found on PATH. Install it (e.g. `brew install pandoc` "
            "on macOS, or see https://pandoc.org/installing.html) to export PDF."
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    # Try a nicer PDF engine first; fall back to default if it's missing.
    candidates = ["xelatex", "pdflatex", "wkhtmltopdf", None]
    last_err: subprocess.CalledProcessError | None = None
    for engine in candidates:
        cmd: list[str] = ["pandoc", str(md_path), "-o", str(path)]
        if engine:
            cmd.extend(["--pdf-engine", engine])
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            return path
        except FileNotFoundError:
            # The engine binary is not installed — try the next.
            continue
        except subprocess.CalledProcessError as e:
            last_err = e
            stderr = (e.stderr or b"").decode("utf-8", errors="replace")
            if "pdf-engine" in stderr or "Could not find" in stderr:
                continue
            raise RuntimeError(
                f"pandoc failed converting to PDF: {stderr.strip() or e}"
            ) from e
    raise RuntimeError(
        "pandoc could not find a working PDF engine. "
        "Install one of: xelatex (TeX Live / MacTeX), pdflatex, or wkhtmltopdf."
        + (f"\nLast error: {last_err}" if last_err else "")
    )


def write_docx(cv_dict: dict[str, Any], path: Path) -> Path:
    """
    Build a clean .docx from the structured optimized CV dict (NOT from the
    Markdown). Keeping it structural avoids relying on pandoc for users who
    just want a Word file.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        ) from e

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Default style tweak — readable serif at 11pt.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    pi = cv_dict.get("personal_info", {})
    name = pi.get("name") or "First Last"
    doc.add_heading(name, level=0)

    title = (cv_dict.get("_offer", {}) or {}).get("position") or pi.get("current_title", "")
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True

    contact_parts = [pi.get(k, "") for k in ("email", "phone", "location", "linkedin", "github", "portfolio")]
    contact_parts = [c for c in contact_parts if c]
    if contact_parts:
        doc.add_paragraph(" · ".join(contact_parts))

    # Summary
    if cv_dict.get("summary"):
        doc.add_heading("Professional summary", level=1)
        doc.add_paragraph(cv_dict["summary"])

    # Experiences
    if cv_dict.get("experiences"):
        doc.add_heading("Professional experience", level=1)
        for exp in cv_dict["experiences"]:
            heading = f"{exp.get('position','')} — {exp.get('company','')}"
            doc.add_heading(heading, level=2)
            period = f"{exp.get('start_date','')} – {exp.get('end_date','')}"
            loc = exp.get("location") or ""
            meta_line = period + (f" · {loc}" if loc else "")
            p = doc.add_paragraph()
            run = p.add_run(meta_line)
            run.italic = True
            for bullet in exp.get("achievements", []):
                doc.add_paragraph(bullet, style="List Bullet")
            techs = exp.get("technologies") or []
            if techs:
                p = doc.add_paragraph()
                p.add_run("Stack: ").bold = True
                p.add_run(", ".join(techs))

    # Skills
    skills = cv_dict.get("skills") or {}
    if skills:
        doc.add_heading("Technical skills", level=1)
        if isinstance(skills, dict):
            for cat, items in skills.items():
                if not items:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{cat}: ").bold = True
                p.add_run(", ".join(items))
        elif isinstance(skills, list):
            doc.add_paragraph(", ".join(skills))

    # Education
    if cv_dict.get("education"):
        doc.add_heading("Education", level=1)
        for ed in cv_dict["education"]:
            line = f"{ed.get('degree','')} — {ed.get('institution','')} ({ed.get('period','')})"
            doc.add_paragraph(line, style="List Bullet")

    # Certifications
    if cv_dict.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in cv_dict["certifications"]:
            issuer = f" — {c.get('issuer','')}" if c.get("issuer") else ""
            year = f" ({c.get('year','')})" if c.get("year") else ""
            doc.add_paragraph(f"{c.get('name','')}{issuer}{year}", style="List Bullet")

    # Languages
    if cv_dict.get("languages"):
        doc.add_heading("Languages", level=1)
        for lang in cv_dict["languages"]:
            doc.add_paragraph(
                f"{lang.get('language','')}: {lang.get('level','')}",
                style="List Bullet",
            )

    doc.save(str(path))
    return path


# ──────────────────────────────────────────────────────────────────────
# Top-level orchestrator
# ──────────────────────────────────────────────────────────────────────
def export_all(
    formats: list[str],
    md_text: str,
    cv_dict: dict[str, Any],
    base_output: Path,
) -> dict[str, Path]:
    """
    Write every requested format. `base_output` is the canonical path
    (e.g. output/cv_optimized.md); other formats reuse the same stem.

    Returns a dict {format: written_path}. Failures for optional formats
    (pdf, docx) are caught and reported but do not abort the others.
    """
    written: dict[str, Path] = {}
    base = base_output.with_suffix("")  # strip .md

    if "md" in formats:
        written["md"] = write_markdown(md_text, base.with_suffix(".md"))

    if "json" in formats:
        written["json"] = write_json(cv_dict, base.with_suffix(".json"))

    if "docx" in formats:
        try:
            written["docx"] = write_docx(cv_dict, base.with_suffix(".docx"))
        except RuntimeError as e:
            written["docx_error"] = Path(str(e))

    if "pdf" in formats:
        # PDF requires the MD on disk.
        md_path = written.get("md") or write_markdown(md_text, base.with_suffix(".md"))
        try:
            written["pdf"] = write_pdf(md_path, base.with_suffix(".pdf"))
        except RuntimeError as e:
            written["pdf_error"] = Path(str(e))

    return written
